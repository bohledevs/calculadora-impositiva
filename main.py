import json
import random
import copy
import login
import csv
from llaves import *
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from facturacion import mostrar_menu_facturas
import os
from facturacion import imprimir_factura
from functools import reduce
from docx import Document

# Tupla de provincias argentinas
provincias_argentina = (    
    "Buenos Aires",
    "Catamarca",
    "Chaco",
    "Chubut",
    "Córdoba",
    "Corrientes",
    "Entre Ríos",
    "Formosa",
    "Jujuy",
    "La Pampa",
    "La Rioja",
    "Mendoza",
    "Misiones",
    "Neuquén",
    "Río Negro",
    "Salta",
    "San Juan",
    "San Luis",
    "Santa Cruz",
    "Santa Fe",
    "Santiago del Estero",
    "Tierra del Fuego",
    "Tucumán"
)

condiciones_iva = (
    "Exento",
    "Responsable Inscripto",
    "Consumidor Final"
)

condiciones_iibb = (
    "Local",
    "Multilateral",
    "No inscripto"
)


# El usuario ingresa los datos requeridos en la pantalla.
# La función encapsula los datos en un diccionario, y lo retorna.
def obtener_entrada():
    
    datos_transaccion = {}

    monto = int(input("Ingrese el monto: "))
    
    while monto < 1:
        monto = int(input("Ingrese un monto correcto: "))
        
    datos_transaccion[LLAVE_MONTO] = monto 

    condicion_fiscal_iva = int(input(f"Cual es su condicion fiscal frente al IVA? \n{imprimir_tupla(condiciones_iva)}\n Respuesta: "))
    
    while condicion_fiscal_iva < 1 or condicion_fiscal_iva > 3:
        condicion_fiscal_iva = int(input(f"Ingrese una condicion fiscal valida: \n{imprimir_tupla(condiciones_iva)}\n Respuesta: "))

    if condicion_fiscal_iva == 1 :
        condicion_fiscal_iva = condiciones_iva[0]
    elif condicion_fiscal_iva == 2 :
        condicion_fiscal_iva = condiciones_iva[1]
    else:
        condicion_fiscal_iva = condiciones_iva[2]
    
    datos_transaccion[LLAVE_CF_IVA] = condicion_fiscal_iva

    condicion_fiscal_iibb = int(input(f"Cual es su condicion fiscal frente a Ingresos Brutos? \n{imprimir_tupla(condiciones_iibb)} \n Respuesta: "))

    while condicion_fiscal_iibb < 1 or condicion_fiscal_iibb > 3:
        condicion_fiscal_iibb = int(input(f"Ingrese una condicion fiscal valida: \n {imprimir_tupla(condiciones_iibb)} \n Respuesta: "))
        
    if condicion_fiscal_iibb == 1 :
        condicion_fiscal_iibb = condiciones_iibb[0]
    elif condicion_fiscal_iibb == 2 :
        condicion_fiscal_iibb = condiciones_iibb[1]
    else:
        condicion_fiscal_iibb = condiciones_iibb[2]
        
    datos_transaccion[LLAVE_CF_IIBB] = condicion_fiscal_iibb
    
    indice_provincia = int(input(f"Ingrese el número correspondiente a su provincia: \n {imprimir_tupla(provincias_argentina)} \n Respuesta: "))
    while indice_provincia < 1 or indice_provincia > 23:
        indice_provincia = int(input(f"Ingrese el número correspondiente a su provincia válido: \n {imprimir_tupla(provincias_argentina)} \n Respuesta: "))
    provincia = provincias_argentina[indice_provincia-1]

    datos_transaccion[LLAVE_PROVINCIA] = provincia

    datos_transaccion[LLAVE_FECHA] = obtener_fecha()

    return datos_transaccion
    

# Imprime tuplas, numeradas por su indice
# Se ingresa una tupla y se imprime en pantalla
def imprimir_tupla(tupla):
    resultado = ""
    for indice, elemento in enumerate(tupla):
        resultado += f"{indice + 1}. {elemento}\n"
    return resultado

# Obtiene la fecha y hora actual en formato 'YYYY-MM-DD_HH-MM-SS'.
def obtener_fecha():
    fecha_actual = datetime.now()
    return fecha_actual.strftime('%Y-%m-%d_%H-%M-%S')

# Recibe los datos del usuario y retorna los impuestos aplicados
# Ingresa un diccionario con los datos de la transacción y retorna una lista con los impuestos aplicados
def calcular_impuestos(datos_transaccion):
    return calcular_nacionales(datos_transaccion) + calcular_provinciales(datos_transaccion)

#Calcula los impuestos nacionales
#Ingresa un diccionario con los datos de la transacción y retorna una lista con los impuestos nacionales
def calcular_nacionales(datos_transaccion):
    condicion_fiscal_iva = datos_transaccion.get(LLAVE_CF_IVA)
    monto = datos_transaccion.get(LLAVE_MONTO)
    return [calcular_iva(monto, condicion_fiscal_iva), calcular_ganancias(monto, condicion_fiscal_iva)]

# Funcion para calcular IVA
#Ingresa el monto de la transacción y la condición fiscal del cliente y retorna un diccionario con el impuesto aplicado
def calcular_iva(monto, cf):
    
    resumen = {}

    tasa = 0
    if (cf == condiciones_iva[1]):
        tasa = 10.5
    elif (cf == condiciones_iva[2]):
        tasa = 21.0
    resumen[LLAVE_TASA] = tasa

    impuesto_aplicado = monto * (tasa/100)
    resumen[LLAVE_IMPUESTO] = impuesto_aplicado
    resumen[LLAVE_TITULO] = "IVA"

    return resumen

# Funcion para calcular IVA
#Ingresa el monto de la transacción y la condición fiscal del cliente y retorna un diccionario con el impuesto aplicado
def calcular_ganancias(monto, cf):
    
    resumen = {}
    resumen[LLAVE_TITULO] = "Ganancias"

    tasa = 0
    if (cf == condiciones_iva[1]):
        tasa = 0.5
    elif (cf == condiciones_iva[2]):
        tasa = 2.0
    resumen[LLAVE_TASA] = tasa

    impuesto_aplicado = monto * (tasa/100)
    resumen[LLAVE_IMPUESTO] = impuesto_aplicado

    return resumen

#Calcula los impuestos provinciales
#Ingresa un diccionario con los datos de la transacción y retorna una lista con los impuestos provinciales
def calcular_provinciales(datos_transaccion):
    condicion_fiscal_iibb = datos_transaccion.get(LLAVE_CF_IIBB)
    monto = datos_transaccion.get(LLAVE_MONTO)
    provincia = datos_transaccion.get(LLAVE_PROVINCIA)
    return [calcular_iibb(monto, condicion_fiscal_iibb, provincia)]

#Calcula el impuesto de Ingresos Brutos
#Ingresa el monto de la transaccion, la condicion fiscal del cliente y la provincia y retorna un diccionario con el impuesto aplicado
def calcular_iibb(monto, cf, provincia):

    resumen = {}
    resumen[LLAVE_TITULO] = "IIBB"

    matriz = []
    
    with open('matriz_iibb.csv', mode='r') as file:
        reader = csv.reader(file)
        for row in reader:
            matriz.append([float(valor) for valor in row])  # Convertir los valores a enteros


    indice_jurisdiccion = provincias_argentina.index(provincia)
    indice_cf = condiciones_iibb.index(cf)
    
    tasa = matriz[indice_jurisdiccion][indice_cf]
    resumen[LLAVE_TASA] = tasa

    impuesto_aplicado = monto * (tasa/100)
    resumen[LLAVE_IMPUESTO] = impuesto_aplicado

    return resumen
  
# Amalgama los datos de la transaccion en una sola salida
def obtener_resumen(datos_transaccion, impuestos_aplicados, usuario):
    resumen = datos_transaccion.copy() 
    resumen[IMPUESTOS_APLICADOS] = impuestos_aplicados
    resumen[IMPUESTOS_AGREGADOS] = impuestos_agregados(impuestos_aplicados)
    resumen[LLAVE_USUARIO] = usuario
    return resumen 


#Funcion monto agregado
def agregar_monto(datos): 
    #Se filtran elementos donde "impuesto" sea igual a 0
    resultados_filtrados = list(filter(lambda x: x["impuesto"] == 0, datos))
    #Usamos map para extraer solo el valor de "impuesto" de los resultados filtrados
    impuestos_filtrados = list(map(lambda x: x["impuesto"], resultados_filtrados))
    #Usamos reduce para sumar todos los valores del campo "impuesto" en la lista original
    total_impuestos = reduce(lambda acc, x: acc + x["impuesto"], datos, 0)
    return total_impuestos

#Funcion para agregar impuestos
def impuestos_agregados(impuestos_aplicados):
    total = 0
    for impuesto in impuestos_aplicados:
        total += impuesto[LLAVE_IMPUESTO]
    return total

#Funcion para imprimir el resumen de una transaccion.
def imprimir_resumen(resumen):
    # Extraemos los datos de la transacción
    monto = resumen['monto']
    condicion_iva = resumen['condicion_fiscal_iva']
    condicion_iibb = resumen['condicion_fiscal_iibb']
    provincia = resumen['provincia']

    # Título principal
    print("==== Resumen de la Transacción ====\n")
    
    # Detalles de la transacción
    print(f"Provincia: {provincia}")
    print(f"Monto: ${monto:.2f}")
    print(f"Condición Fiscal IVA: {condicion_iva}")
    print(f"Condición Fiscal IIBB: {condicion_iibb}\n")
    
    # Título para los impuestos
    print("==== Impuestos Aplicados ====\n")

    # Detalles de los impuestos aplicados
    for impuesto in resumen[IMPUESTOS_APLICADOS]:
        titulo = impuesto['titulo']
        tasa = impuesto['tasa']
        monto_impuesto = impuesto['impuesto']
        print(f"{titulo}:")
        print(f"  - Tasa: {tasa}%")
        print(f"  - Monto del impuesto: ${monto_impuesto:.2f}\n")
        
# Funcion para crear carpeta con los errores si no existe
def crear_carpeta_errores(nombre_carpeta):
    if not os.path.exists(nombre_carpeta):
        os.makedirs(nombre_carpeta)
        
# Funcion para crear PDF de error
def guardar_error(id_error, fecha, usuario, descripcion_error, filename):
    nombre_carpeta = "errores"
    crear_carpeta_errores(nombre_carpeta)
    
    archivo_errores = os.path.join(nombre_carpeta, "errores.docx")

     # Crear o abrir el archivo de errores
    if os.path.exists(archivo_errores):
        # Si el archivo existe, abrirlo
        documento = Document(archivo_errores)
    else:
        # Si no existe, crear uno nuevo
        documento = Document()
        documento.add_heading("Registro de Errores", level=1)

    # Agregar un nuevo error al archivo
    documento.add_paragraph("=" * 50)
    documento.add_paragraph(f"ID de Error: {id_error}")
    documento.add_paragraph(f"Fecha: {fecha}")
    documento.add_paragraph(f"Usuario: {usuario}")
    documento.add_paragraph(f"Descripción del Error: {descripcion_error}")
    documento.add_paragraph("-" * 50)

    # Guardar el archivo
    documento.save(archivo_errores)
   
    print(f"Error registrado en el archivo {filename}")
    

 
def mostrar_banner(titulo):
    # Imprime una línea de 50 signos de igual para crear una separación visual
    print("=" * 50)
     # Imprime el título centrado con un margen de 10 espacios en blanco a la izquierda
    print(" " * 10 + titulo)
     # Imprime otra línea de 50 signos de igual para cerrar la separación visual
    print("=" * 50)
 
def mostrar_menu_principal():
    # Define una tupla con las opciones del menú principal
    opciones = (
    "Calcular impuestos",
    "Ver mis facturas",
    "Salir"
    )

     # Muestra un banner con el título "Calculadora Impositiva"
    mostrar_banner("Calculadora Impositiva")
    # Imprime las opciones del menú
    print(imprimir_tupla(opciones))
    # Imprime una línea de 50 signos de igual para crear una separación visual
    print("=" * 50)
    # Solicita al usuario que ingrese una opción y la retorna como un entero
    return int(input("Ingrese su opción: "))


## PROGRAMA PRINCIPAL
def programa_principal():
    try:
        usuario = login.menu()
        print(usuario)
        if not usuario:
            return
        op = mostrar_menu_principal()
        if (1 == op):
            datos_transaccion = obtener_entrada()
            impuestos_aplicados = calcular_impuestos(datos_transaccion)
            resumen = obtener_resumen(datos_transaccion, impuestos_aplicados, usuario)
            imprimir_resumen(resumen)
            imprimir_factura(resumen)
        elif (2 == op):
            mostrar_menu_facturas()
        elif (3 == op):
            print("Finalizando programa.")
        else:
            raise IndexError(f"No eligió una opción válida:{op}")
    except Exception as e:
        # Obtener detalles del error
        id_error = random.randint(1000, 9999)  # ID de error aleatorio
        fecha_error = obtener_fecha()
        usuario = usuario["nombre"]  # Aquí deberías agregar la variable del usuario actual, si es posible.
        descripcion_error = str(e)  # Convertir el error a una cadena
        filename = f"error_{id_error}"
        
        
        # Crear el PDF con los detalles del error
        guardar_error(id_error, fecha_error, usuario, descripcion_error,filename )

        print("Se generó un archivo con los detalles del error.")

programa_principal()
